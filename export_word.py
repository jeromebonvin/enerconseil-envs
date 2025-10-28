"""
Export des données JSON extraites vers un document Word formaté (Document de Conformité)
Version CORRIGÉE avec remplissage des tableaux et corrections orthographiques
Usage : python export_word.py fichier_extraction.json
"""

import json
import argparse
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from utils import (
    set_font,
    remove_table_borders,
    set_table_header_gray,
    info,
    success,
    warning,
    error,
)


# -----------------------------------------------------
# CORRECTION ORTHOGRAPHIQUE
# -----------------------------------------------------
def corriger_orthographe_champ(nom_champ):
    """Corrige l'orthographe et la casse des noms de champs"""
    corrections = {
        # Générales
        'sre': 'SRE',
        'snf': 'SNF',
        'sre neuf': 'SRE neuf',
        'sre existant': 'SRE existant',
        'egid': 'EGID',
        'ecs': 'ECS',
        'pac': 'PAC',
        'pc': 'PC',
        'sia': 'SIA',
        'minergie': 'Minergie',
        'ehwlk': 'EHWLK',
        'qh': 'Qh',
        
        # EN-VS
        'type de batiment': 'Type de bâtiment',
        'type batiment': 'Type de bâtiment',
        'type de construction': None,  # À supprimer
        'affectation principale': 'Affectation principale',
        'nombre d etages': "Nombre d'étages",
        'nombre etages': "Nombre d'étages",
        'annee construction': 'Année construction',
        'annee de construction': 'Année de construction',
        'numero parcelle': 'Numéro parcelle',
        'maitre ouvrage': "Maître d'ouvrage",
        'maitre d ouvrage': "Maître d'ouvrage",
        'nature des travaux': 'Nature des travaux',
        'projet interet cantonal': None,  # ✅ À supprimer selon demande
        "projet d'interet cantonal": None,  # ✅ À supprimer selon demande
        'categorie sia': 'Catégorie SIA',
        'type de chauffage': 'Type de chauffage',
        'couverture des besoins de chaleur': 'Couverture des besoins de chaleur',
        'formulaires necessaires ou annexes': 'Formulaires annexés',
        
        # EN-VS-101a
        'rafraichissement humidification ou deshumidification': 'Rafraîchissement et/ou (dés)humidification',
        'pac reversible projetee': 'PAC réversible projetée',
        'solution standard choisie': 'Solution standard choisie',
        
        # EN-VS-102a
        'hygiene air interieur concept ventilation': 'Concept de ventilation',
        'protection thermique ete valeur g': 'Respect de la valeur g',
        'protection thermique ete rafraichissement': 'Rafraîchissement',
        'valeurs u respectees par tous elements': 'Valeurs U respectées par tous les éléments',
        'justificatif ponts thermiques respecte': 'Justificatif des ponts thermiques respecté',
        'enveloppe thermique completement fermee': 'Enveloppe thermique complètement fermée',
        'tous locaux chauffes interieur enveloppe': "Tous les locaux chauffés à l'intérieur de l'enveloppe thermique",
        
        # EN-VS-103
        'type de generateur de chaleur': 'Type de générateur de chaleur',
        'puissance thermique kw': 'Puissance thermique',
        'dispositif comptage energie': "Dispositif de comptage d'énergie",
        'emission chaleur locaux isoles': 'Émission de chaleur que dans locaux isolés',
        'emission de chaleur': 'Émission de chaleur',
        'temperature emission chaleur': "Température de départ de l'émetteur de chaleur",
        'regulation temperature par local': 'Régulation de température par local',
        'temperature ecs inferieure 60': 'Température ECS < 60°C',
        'systeme mesure installe chauffage': 'Système de mesure installé pour chauffage',
        'systeme mesure installe ecs': 'Système de mesure installé pour ECS',
        
        # EN-VS-104
        'puissance production propre electricite annuelle': 'Puissance projetée',
        'puissance production propre electricite requise': 'Puissance requise',
        'nbre de panneaux': 'Nombre de panneaux',
        'punitaire des panneaux wc': 'Puissance unitaire des panneaux',
        
        # EN-VS-110
        'surface nette plancher rafraichi deshumidifie': 'Surface nette de plancher rafraîchie',
        'total puissances thermiques frigorifiques': 'Total des puissances frigorifiques',
        'puissances electriques total kw': 'Puissance électrique totale',
        'puissances electriques puissance specifique w m2': 'Puissance spécifique',
        
        # Autres
        'responsable': 'Responsable',
        'architecte': 'Architecte',
        'telephone': 'Téléphone',
        'thermique': 'Thermique',
        'energie': 'Énergie',
        'batiment': 'Bâtiment',
        'electrique': 'Électrique',
        'electricite': 'Électricité',
        'categorie': 'Catégorie',
        'renouvellement': 'Renouvellement',
        'renovation': 'Rénovation',
        'chauffage': 'Chauffage',
        'ventilation': 'Ventilation',
        'refroidissement': 'Refroidissement',
        'climatisation': 'Climatisation',
        'eau chaude sanitaire': 'Eau chaude sanitaire',
        'enveloppe': 'Enveloppe',
        'isolation': 'Isolation',
        'coefficient': 'Coefficient',
        'puissance': 'Puissance',
        'rendement': 'Rendement',
        'surface': 'Surface',
        'volume': 'Volume',
        'deperdition': 'Déperdition',
        'apport': 'Apport',
        'besoin': 'Besoin',
        'production': 'Production',
        'consommation': 'Consommation',
    }
    
    nom_lower = nom_champ.lower()
    
    # Chercher une correspondance exacte
    if nom_lower in corrections:
        result = corrections[nom_lower]
        if result is None:  # Champ à supprimer
            return None
        return result
    
    # Chercher une correspondance partielle
    for pattern, replacement in corrections.items():
        if replacement is not None and pattern in nom_lower:
            nom_champ = nom_champ.lower().replace(pattern, replacement)
    
    # Capitaliser la première lettre si ce n'est pas déjà fait
    if nom_champ and not nom_champ[0].isupper():
        nom_champ = nom_champ[0].upper() + nom_champ[1:]
    
    return nom_champ


# -----------------------------------------------------
# 1️⃣ EN-TÊTE DU DOCUMENT
# -----------------------------------------------------
def ajouter_entete(doc, numero_dossier=None, config=None):
    """Ajoute l'en-tête (logo, adresse, date, introduction)."""
    header_table = doc.add_table(rows=3, cols=2)
    remove_table_borders(header_table)

    # Logo + adresse
    cell_logo = header_table.rows[0].cells[0]
    cell_logo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_path = "logo.png"
    if os.path.exists(logo_path):
        try:
            run = cell_logo.paragraphs[0].add_run()
            run.add_picture(logo_path, width=Inches(2.25))
        except Exception as e:
            warning(f"Logo non inséré : {e}")
    p = cell_logo.add_paragraph(
        "Passage de la Matze 13\nCH-1950 Sion\n027 558 83 47\nwww.enerconseil.ch\nsion@enerconseil.ch"
    )
    p.paragraph_format.space_before = Pt(0)
    set_font(p)

    # Destinataire (depuis config si disponible)
    cell_dest = header_table.rows[1].cells[1]
    p = cell_dest.paragraphs[0]
    
    if config and 'commune' in config:
        commune_info = config['commune']
        p.text = f"{commune_info['nom']}\n{commune_info['adresse']}"
    else:
        # ✅ Pas de valeur par défaut si config manquante
        p.text = ""
    
    p.paragraph_format.space_after = Pt(24)
    set_font(p)

    # Date
    mois_fr = {
        "January": "janvier", "February": "février", "March": "mars",
        "April": "avril", "May": "mai", "June": "juin",
        "July": "juillet", "August": "août", "September": "septembre",
        "October": "octobre", "November": "novembre", "December": "décembre"
    }
    now = datetime.now()
    mois_en = now.strftime("%B")
    date_fr = now.strftime(f"%d {mois_fr.get(mois_en, mois_en)} %Y")
    
    cell_date = header_table.rows[2].cells[1]
    p = cell_date.paragraphs[0]
    p.text = f"Sion, le {date_fr}"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(p)

    # Titre et introduction
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Attestation de conformité énergétique").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p, size=14)

    doc.add_paragraph()
    if numero_dossier:
        p = doc.add_paragraph(f"Réf. : Dossier N° {numero_dossier}")
    else:
        p = doc.add_paragraph("Réf. : Dossier N° ")
    set_font(p)

    doc.add_paragraph()
    p = doc.add_paragraph(
        "Madame, Monsieur,\n\n"
        "Suite à l'examen du dossier susmentionné, nous vous confirmons que "
        "votre projet est conforme à la législation cantonale en vigueur "
        "(LcEne, norme SIA 380/1 : 2016 et norme SIA 180)."
    )
    set_font(p)
    doc.add_paragraph()


# -----------------------------------------------------
# 2️⃣ TABLEAU ÉLÉMENTS ENVELOPPE (OPAQUES)
# -----------------------------------------------------
def creer_tableau_elements_enveloppe_opaque(doc, elements_opaques):
    """Crée un tableau pour les éléments d'enveloppe opaques."""
    if not elements_opaques:
        return

    p = doc.add_paragraph()
    run = p.add_run("Tableau récapitulatif éléments d'enveloppe")
    run.bold = True
    set_font(p, size=12)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_header_gray(table, row_index=0)

    headers = ["Élément", "U [W/m²K]", "Surface [m²]", "Remarque"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        set_font(cell.paragraphs[0], bold=True)

    for elem in elements_opaques:
        row = table.add_row()
        row.cells[0].text = elem.get("nom", "")
        row.cells[1].text = str(elem.get("u", ""))
        row.cells[2].text = str(elem.get("surface", ""))
        row.cells[3].text = elem.get("remarque", "")

        for cell in row.cells:
            set_font(cell.paragraphs[0])

    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(1.5)
    table.columns[3].width = Inches(1.5)


def ajouter_elements_vitres_au_tableau(table, elements_vitres):
    """Ajoute des éléments vitrés au tableau existant en regroupant les fenêtres identiques."""
    
    # ✅ Regrouper les fenêtres avec les mêmes propriétés (Ug, g, Uf)
    fenetres_groupees = {}
    autres_elements = []
    
    for elem in elements_vitres:
        nom = elem.get("nom", "").lower()
        
        # Si c'est une fenêtre, on regroupe par propriétés
        if "fenetre" in nom or "fenêtre" in nom:
            ug = elem.get("ug", "")
            g = elem.get("g", "")
            uf = elem.get("uf", "")
            
            # Créer une clé unique basée sur les propriétés
            cle = (ug, g, uf)
            
            if cle not in fenetres_groupees:
                fenetres_groupees[cle] = {
                    "ug": ug,
                    "g": g,
                    "uf": uf,
                    "surface": 0,
                    "remarque": elem.get("remarque", "")
                }
            
            # Ajouter la surface
            try:
                surface = float(elem.get("surface", 0))
                fenetres_groupees[cle]["surface"] += surface
            except (ValueError, TypeError):
                pass
        else:
            # Autres éléments vitrés (portes, etc.)
            autres_elements.append(elem)
    
    # Ajouter les fenêtres regroupées
    for proprietes, donnees in fenetres_groupees.items():
        row = table.add_row()
        row.cells[0].text = "Fenêtre"  # ✅ Nom générique
        
        # Construire le texte pour U/g/Uf
        info_vitrage = []
        if donnees["ug"]:
            info_vitrage.append(f"Ug={donnees['ug']}")
        if donnees["g"]:
            info_vitrage.append(f"g={donnees['g']}")
        if donnees["uf"]:
            info_vitrage.append(f"Uf={donnees['uf']}")
        
        row.cells[1].text = "\n".join(info_vitrage)
        row.cells[2].text = f"{donnees['surface']:.2f}" if donnees['surface'] else ""
        row.cells[3].text = donnees.get("remarque", "")
        
        for cell in row.cells:
            set_font(cell.paragraphs[0])
    
    # Ajouter les autres éléments vitrés
    for elem in autres_elements:
        row = table.add_row()
        row.cells[0].text = elem.get("nom", "")
        
        info_vitrage = []
        if elem.get("ug"):
            info_vitrage.append(f"Ug={elem.get('ug')}")
        if elem.get("g"):
            info_vitrage.append(f"g={elem.get('g')}")
        if elem.get("uf"):
            info_vitrage.append(f"Uf={elem.get('uf')}")
        
        row.cells[1].text = "\n".join(info_vitrage)
        row.cells[2].text = str(elem.get("surface", ""))
        row.cells[3].text = elem.get("remarque", "")
        
        for cell in row.cells:
            set_font(cell.paragraphs[0])


# -----------------------------------------------------
# 3️⃣ TABLEAU FORMULAIRE
# -----------------------------------------------------
def creer_tableau_formulaire(doc, titre, donnees, form_id):
    """Crée un tableau pour un formulaire donné."""
    p = doc.add_paragraph()
    run = p.add_run(titre)
    run.bold = True
    set_font(p, size=12)

    elements_opaques = donnees.pop("elements_opaques", [])
    elements_vitres = donnees.pop("elements_vitres", [])

    if donnees:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        set_table_header_gray(table, row_index=0)

        # En-tête
        cell = table.rows[0].cells[0]
        cell.text = "Champ"
        set_font(cell.paragraphs[0], bold=True)

        cell = table.rows[0].cells[1]
        cell.text = "Valeur"
        set_font(cell.paragraphs[0], bold=True)

        # ✅ Traitement spécial pour EN-VS-104 : regrouper les infos panneaux
        if form_id == 'EN-VS-104':
            nbre_panneaux = None
            puissance_unitaire = None
            donnees_temp = {}
            
            for cle, valeur in donnees.items():
                cle_lower = cle.lower()
                if 'nbre' in cle_lower and 'panneaux' in cle_lower:
                    nbre_panneaux = valeur
                elif 'punitaire' in cle_lower or ('puissance' in cle_lower and 'unitaire' in cle_lower):
                    puissance_unitaire = valeur
                else:
                    donnees_temp[cle] = valeur
            
            # Ajouter la ligne regroupée si les deux valeurs sont présentes
            if nbre_panneaux is not None and puissance_unitaire is not None:
                donnees_temp['Installation prévue'] = f"{nbre_panneaux} panneaux de {puissance_unitaire} Wc"
            
            donnees = donnees_temp
        
        # ✅ Traitement spécial pour EN-VS-101b : renommer et ordonner les champs
        if form_id == 'EN-VS-101b':
            donnees_ordonnees = {}
            ehwlk_limite = None
            ehwlk_calculee = None
            limite_respectee = None
            
            for cle, valeur in donnees.items():
                cle_lower = cle.lower()
                if 'ehwlk' in cle_lower and 'limite' in cle_lower and 'respecte' not in cle_lower:
                    ehwlk_limite = valeur
                elif 'ehwlk' in cle_lower and 'calcul' in cle_lower:
                    ehwlk_calculee = valeur
                elif 'limite' in cle_lower and 'respecte' in cle_lower:
                    limite_respectee = valeur
                else:
                    donnees_ordonnees[cle] = valeur
            
            # Ajouter dans l'ordre souhaité
            if ehwlk_limite is not None:
                donnees_ordonnees['Valeur EHWLK limite'] = ehwlk_limite
            if ehwlk_calculee is not None:
                donnees_ordonnees['Valeur EHWLK calculée'] = ehwlk_calculee
            if limite_respectee is not None:
                donnees_ordonnees['Valeur limite respectée'] = limite_respectee
            
            donnees = donnees_ordonnees
        
        # ✅ Traitement spécial pour EN-VS-102b : renommer et ordonner les champs avec unités
        if form_id == 'EN-VS-102b':
            donnees_ordonnees = {}
            qh = None
            qh_li = None
            puissance_spec = None
            puissance_limite = None
            
            for cle, valeur in donnees.items():
                cle_lower = cle.lower()
                if 'qh' in cle_lower and 'limite' not in cle_lower and 'li' not in cle_lower:
                    qh = valeur
                elif 'qh' in cle_lower and ('limite' in cle_lower or 'li' in cle_lower):
                    qh_li = valeur
                elif 'puissance' in cle_lower and 'specifique' in cle_lower:
                    puissance_spec = valeur
                elif 'puissance' in cle_lower and 'limite' in cle_lower:
                    puissance_limite = valeur
                else:
                    donnees_ordonnees[cle] = valeur
            
            # Ajouter dans l'ordre souhaité avec unités dans la colonne 2
            if qh is not None:
                donnees_ordonnees['Besoins de chaleur pour le chauffage Qh'] = f"{qh} kWh/m²"
            if qh_li is not None:
                donnees_ordonnees['Besoins de chaleur limite Qh,li'] = f"{qh_li} kWh/m²"
            if puissance_spec is not None:
                donnees_ordonnees['Puissance de chauffage spécifique'] = f"{puissance_spec} W/m²"
            if puissance_limite is not None:
                donnees_ordonnees['Puissance de chauffage limite'] = f"{puissance_limite} W/m²"
            
            donnees = donnees_ordonnees

        # Contenu
        for cle, valeur in donnees.items():
            # ✅ Correction orthographique et suppression des champs à retirer
            cle_corrigee = corriger_orthographe_champ(cle)
            
            if cle_corrigee is None:  # Champ à supprimer
                continue
            
            # ✅ Ne pas afficher si la valeur est vide, None ou valeur par défaut
            if valeur in [None, "", "N/A", "Non spécifié", "À compléter"]:
                continue
            
            row_cells = table.add_row().cells
            row_cells[0].text = cle_corrigee
            set_font(row_cells[0].paragraphs[0])

            # Formatage valeur
            if isinstance(valeur, bool):
                row_cells[1].text = "Oui" if valeur else "Non"
            elif isinstance(valeur, (int, float)):
                row_cells[1].text = str(valeur)
            else:
                row_cells[1].text = str(valeur)
            set_font(row_cells[1].paragraphs[0])
            
            # ✅ Surligner en jaune la ligne "Solution standard choisie"
            if form_id == 'EN-VS-101a' and 'solution' in cle.lower() and 'standard' in cle.lower():
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                for cell in row_cells:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))  # Jaune
                    cell._element.get_or_add_tcPr().append(shading_elm)
            
            # ✅ Ajouter la ligne "Commande automatique des protections solaires" après "Respect de la valeur g"
            if (form_id in ['EN-VS-102a', 'EN-VS-102b']) and \
               ('protection thermique ete valeur g' in cle.lower() or 'respect de la valeur g' in cle.lower()):
                row_cmd = table.add_row().cells
                row_cmd[0].text = "Commande automatique des protections solaires"
                row_cmd[1].text = ""  # ✅ Valeur vide à remplir
                set_font(row_cmd[0].paragraphs[0])
                set_font(row_cmd[1].paragraphs[0])

        table.columns[0].width = Inches(3.0)
        table.columns[1].width = Inches(4.0)
        doc.add_paragraph()

    # Tableaux d'enveloppe
    if elements_opaques or elements_vitres:
        creer_tableau_elements_enveloppe_opaque(doc, elements_opaques)
        
        # Ajouter les éléments vitrés dans le même tableau
        if elements_vitres:
            # Retrouver le dernier tableau créé
            tables = doc.tables
            if tables:
                dernier_tableau = tables[-1]
                ajouter_elements_vitres_au_tableau(dernier_tableau, elements_vitres)
        
        doc.add_paragraph()


# -----------------------------------------------------
# 4️⃣ PIED DE PAGE
# -----------------------------------------------------
def ajouter_pied_de_page(doc, config=None):
    doc.add_paragraph()

    # Remarques
    remarques_table = doc.add_table(rows=1, cols=1)
    remarques_table.style = "Table Grid"
    cell = remarques_table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.text = "Remarques :"
    set_font(p, bold=True)

    remarques = [
        "Le bâtiment doit être équipé d'un système de mesure par unité d'occupation "
        "pour le chauffage et l'eau chaude sanitaire (décompte individuel des frais "
        "de chauffage et d'ECS).",

        "Étant donné les hauteurs de fenêtres, nous vous recommandons fortement d'installer "
        "du triple vitrage, selon les indications de la norme SIA 180:2014 (point 4.1.3), "
        "sans quoi les exigences de confort hivernal dans la cuisine et séjour ne sont pas respectées.",

        "Nous recommandons l'appui d'un physicien du bâtiment pour éviter tout risque de "
        "condensation et de moisissures au niveau des ponts thermiques et autres raccords."
    ]
    for r in remarques:
        p = cell.add_paragraph(r, style="List Bullet")
        set_font(p)

    # Préavis
    doc.add_paragraph()
    preavis_table = doc.add_table(rows=1, cols=1)
    preavis_table.style = "Table Grid"
    
    # ✅ Encadrer en vert
    tbl = preavis_table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # Épaisseur
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '00B050')  # Vert
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    cell = preavis_table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.text = "Préavis"
    set_font(p, bold=True)

    preavis = [
        "Le projet de construction présenté respecte les exigences légales "
        "(LcEne, SIA 380/1 et SIA 180), conformément au document de conformité "
        "joint au préavis.",

        "Le projet respecte les exigences légales (normes SIA 380/1 : 2016 et SIA 180), "
        "sous réserve de la validation du justificatif EN-VS-105 qui sera envoyé plus tard. "
        "Ce dernier doit nous parvenir pour contrôle au plus tard 8 semaines avant le début "
        "des travaux, comme mentionné sur le formulaire EN-VS."
    ]
    for r in preavis:
        p = cell.add_paragraph(r, style="List Bullet")
        set_font(p)
    
    doc.add_paragraph()

    # Texte final
    p = doc.add_paragraph(
        "Les indications figurant dans ce rapport doivent être respectées. "
        "Des contrôles sur chantier peuvent être effectués."
    )
    set_font(p)

    p = doc.add_paragraph(
        "En vous souhaitant bonne réception, nous vous prions d'agréer, "
        "Mesdames, Messieurs, nos meilleures salutations."
    )
    set_font(p)

    doc.add_paragraph()

    # Signature (depuis config si disponible)
    sig_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Cm(10)
    sig_table.columns[1].width = Cm(7)
    
    cell = sig_table.rows[0].cells[1]
    p = cell.paragraphs[0]
    
    if config and 'utilisateur' in config:
        utilisateur_info = config['utilisateur']
        p.text = f"Enerconseil SA\n{utilisateur_info['prenom']} {utilisateur_info['nom']}"
    else:
        # ✅ Pas de valeur par défaut si config manquante
        p.text = "Enerconseil SA\n"
    
    set_font(p)


# -----------------------------------------------------
# 5️⃣ FONCTIONS UTILITAIRES POUR VÉRIFIER LES FORMULAIRES
# -----------------------------------------------------
def verifier_formulaire_rempli(donnees, form_id):
    """
    Vérifie si un formulaire est présent et contient des données significatives.
    Retourne True si le formulaire est rempli, False sinon.
    """
    if form_id not in donnees:
        return False
    
    form_data = donnees[form_id]
    
    # Ignorer les clés liées aux éléments d'enveloppe
    cles_a_ignorer = ['elements_opaques', 'elements_vitres']
    
    # Vérifier s'il y a au moins une valeur significative
    for cle, valeur in form_data.items():
        if cle in cles_a_ignorer:
            continue
        
        # Une valeur est significative si elle n'est pas vide/None/valeur par défaut
        if valeur not in [None, "", "N/A", "Non spécifié", "À compléter"]:
            return True
    
    return False


# -----------------------------------------------------
# 6️⃣ EXPORT PRINCIPAL
# -----------------------------------------------------
def exporter_vers_word(json_path, output_path=None):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Extraire la configuration si présente
    config = data.pop('_config', None)

    doc = Document()
    ajouter_entete(doc, data.get("numero_dossier"), config)

    # ✅ Vérifier quels formulaires 101 et 102 sont présents et remplis
    form_101a_present = verifier_formulaire_rempli(data, "EN-VS-101a")
    form_101b_present = verifier_formulaire_rempli(data, "EN-VS-101b")
    form_102a_present = verifier_formulaire_rempli(data, "EN-VS-102a")
    form_102b_present = verifier_formulaire_rempli(data, "EN-VS-102b")
    
    # ✅ Supprimer les formulaires non remplis
    if form_101a_present and form_101b_present:
        # Si les deux sont présents, garder celui qui a le plus de données
        info("Les deux formulaires EN-VS-101a et EN-VS-101b sont présents. Conservation du plus complet.")
        if len(data.get("EN-VS-101a", {})) < len(data.get("EN-VS-101b", {})):
            data.pop("EN-VS-101a", None)
            form_101a_present = False
        else:
            data.pop("EN-VS-101b", None)
            form_101b_present = False
    elif not form_101a_present and "EN-VS-101a" in data:
        data.pop("EN-VS-101a", None)
    elif not form_101b_present and "EN-VS-101b" in data:
        data.pop("EN-VS-101b", None)
    
    if form_102a_present and form_102b_present:
        # Si les deux sont présents, garder celui qui a le plus de données
        info("Les deux formulaires EN-VS-102a et EN-VS-102b sont présents. Conservation du plus complet.")
        if len(data.get("EN-VS-102a", {})) < len(data.get("EN-VS-102b", {})):
            data.pop("EN-VS-102a", None)
            form_102a_present = False
        else:
            data.pop("EN-VS-102b", None)
            form_102b_present = False
    elif not form_102a_present and "EN-VS-102a" in data:
        data.pop("EN-VS-102a", None)
    elif not form_102b_present and "EN-VS-102b" in data:
        data.pop("EN-VS-102b", None)

    ordre = [
        ("EN-VS", "EN-VS – Informations générales"),
        ("EN-VS-101a", "EN-VS-101a – Couverture des besoins de chaleur - solution standard"),
        ("EN-VS-101b", "EN-VS-101b – Couverture des besoins de chaleur calculée"),
        ("EN-VS-102a", "EN-VS-102a – Protection thermique - exigences ponctuelles"),
        ("EN-VS-102b", "EN-VS-102b – Protection thermique - exigences globales"),
        ("EN-VS-103", "EN-VS-103 – Chauffage et eau chaude sanitaire"),
        ("EN-VS-104", "EN-VS-104 – Production propre d'électricité"),
        ("EN-VS-105", "EN-VS-105 – Installation de ventilation"),
        ("EN-VS-110", "EN-VS-110 – Rafraîchissement, (dés)humidification"),
    ]

    for form_id, titre in ordre:
        if form_id in data:
            info(f"Ajout du formulaire {form_id}")
            creer_tableau_formulaire(doc, titre, data[form_id], form_id)

    ajouter_pied_de_page(doc, config)

    if output_path is None:
        output_path = json_path.replace(".json", "_rapport_conformite.docx")

    try:
        doc.save(output_path)
        success(f"Document Word créé : {output_path}")
    except PermissionError:
        import time
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        output_path = f"{output_path.replace('.docx', '')}_{timestamp}.docx"
        doc.save(output_path)
        warning(f"Fichier Word ouvert, enregistré sous {output_path}")

    return output_path


# -----------------------------------------------------
# 7️⃣ MODE CLI
# -----------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Exporter un fichier JSON d'extraction vers Word (format conformité)")
    parser.add_argument("json_path", help="Chemin vers le JSON d'extraction")
    parser.add_argument("--output", "-o", help="Nom du fichier Word de sortie")
    args = parser.parse_args()

    if not os.path.exists(args.json_path):
        error(f"Fichier JSON non trouvé : {args.json_path}")
        return

    exporter_vers_word(args.json_path, args.output)


if __name__ == "__main__":
    main()
